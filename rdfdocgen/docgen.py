from rdflib import Literal, URIRef, BNode, ConjunctiveGraph
import rdflib
from docx import Document
from re import compile
import pandas as pd
import click
import logging
import os

logging.basicConfig(format='%(asctime)s %(message)s', level=logging.INFO)
logger = logging.getLogger('docgen')


pat = compile('\\W')


def populate_conjunctive_graph(directory):
    g = ConjunctiveGraph()
    for root, dirs, files in os.walk(directory):
        for f in files:
            if f.split('.')[-1] not in rdflib.util.SUFFIX_FORMAT_MAP:
                continue
            fspec = os.path.join(root, f)
            try:
                logger.info('Parsing %s', fspec)
                g.parse(fspec, format=rdflib.util.guess_format(fspec))
            except rdflib.plugin.PluginException:
                continue
    return g


name_table = dict()   # Glossary entries


def try_qname(g, o):
    try:
        name_table.update({g.qname(o): o})
        return g.qname(o)
    except ValueError:
        return o


def process_blank_node(g, bn, doc):
    for tups in g.predicate_objects(bn):
        doc.add_heading(g.qname(tups[0]), level=3)
        doc.add_paragraph(try_qname(g, tups[1]))
        if type(tups[1]) == BNode:
            process_blank_node(g, tups[1], doc)


@click.command()
@click.option('--directory', help='Source RDF directory')
@click.option('--output', help='Output MS Word File')
@click.option('--title', default='RDF Documentation', help='Document Title')
def do_work(directory, output, title):
    logger.info('Input %s', directory)
    logger.info('Output %s', output)
    logger.info('Title %s' % title)
    g = populate_conjunctive_graph(directory)
    glist = [(graph.identifier, graph) for graph in g.contexts()]
    gdf = pd.DataFrame(glist, columns=['id', 'g']).sort_values(['id']).set_index(['id'])
    document = Document()
    document.add_heading(title, level=0)
    for i, gr in gdf.iterrows():
        docnm = str(i).split('/')[-1].split('.')[0]
        logger.info('Processing: %s' % docnm)

        spdf = pd.DataFrame([(s, p, o) for s, p, o in gr.g],
                            columns=['s', 'p', 'o']).sort_values(['s', 'p']).set_index(['s', 'p'])
        for subject, new_df in spdf.groupby(level=0):
            if isinstance(subject, BNode):
                continue
            document.add_heading(try_qname(g, subject), level=1)
            header_text = ''
            for sp, objs in new_df.iterrows():
                pname = try_qname(g, sp[1])
                if header_text != pname:
                    document.add_heading(pname, level=2)
                    header_text = pname
                for o in objs:
                    if type(o) == Literal:
                        document.add_paragraph(o)
                    elif type(o) == URIRef:
                        document.add_paragraph(try_qname(g, o))
                    elif type(o) == BNode:
                        process_blank_node(gr.iloc[0], o, document)
                        # process_blank_node(gr, o, document)
    # Create the glossary
    ntdf = pd.DataFrame([(k, name_table[k]) for k in name_table],
                        columns=['Qualified Name',
                                 'IRI']).sort_values(['Qualified Name',
                                                      'IRI']).set_index(['Qualified Name'])
    document.add_heading('Glossary', level=1)
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qualified Name'
    hdr_cells[1].text = 'IRI'
    for i, r in ntdf.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = r.IRI

    logger.info('Saving %s' % output)
    document.save(output)
    logger.info('Done')


if __name__ == '__main__':
    do_work()
